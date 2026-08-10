# -*- coding: utf-8 -*-
"""
Module sinh văn bản Word: Quyết định nâng lương, Tờ trình Ban Tổ chức Tỉnh ủy,
Biên bản họp Hội đồng xét nâng lương — theo đúng mẫu QDNLTX / TT_NANG_LUONG.

Font: Times New Roman, cỡ 14 (thân bài), 14-bold (tiêu đề mục), căn lề:
Top 2cm, Bottom 2cm, Left 3cm, Right 1.5cm — khổ A4.
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────
# THÀNH PHẦN HỘI ĐỒNG MẶC ĐỊNH (có thể chỉnh sửa trên giao diện)
# ─────────────────────────────────────────────────────────
THANH_PHAN_HOI_DONG_MAC_DINH = """Đồng chí Trần Mạnh Lợi, Ủy viên Ban Thường vụ, Trưởng Ban – Chủ tịch Hội đồng xét nâng lương Ban Tuyên giáo và Dân vận Tỉnh ủy – Chủ trì.
Đồng chí Nguyễn Lam Sơn, Tỉnh ủy viên, Phó Trưởng ban Thường trực.
Đồng chí Lê Mạnh Cường, Phó Trưởng ban.
Đồng chí Nguyễn Văn Hưng, Phó Trưởng ban.
Đồng chí Vương Thúy Hằng, Phó Trưởng ban.
Đồng chí Hoàng Thị Hằng, Phó Trưởng ban.
Đồng chí Chẩu Thị Thu, Phó Trưởng ban.
Đồng chí Đặng Ái Xoan, Phó Trưởng ban.
Đồng chí Nguyễn Thu Vân, Trưởng phòng Tuyên truyền, Báo chí - Xuất bản.
Đồng chí Phan Thanh Bình, Trưởng phòng Khoa giáo, Văn hóa - Văn nghệ.
Đồng chí Trần Văn Mạnh, Trưởng phòng Đoàn thể và các Hội.
Đồng chí Trần Thị Thanh Huyền, Trưởng phòng Dân vận các cơ quan nhà nước, Dân tộc và Tôn giáo.
Đồng chí Đinh Thị Thúy, Chánh văn phòng Ban - Thư ký"""

CAN_CU_THUONG_XUYEN = "Căn cứ Thông tư số 08/2013/TT-BNV ngày 31/7/2013 của Bộ Nội vụ về Hướng dẫn thực hiện chế độ nâng lương thường xuyên và nâng lương trước thời hạn đối với cán bộ, công chức, viên chức và người lao động;"
CAN_CU_VUOT_KHUNG = "Căn cứ Thông tư số 03/2021/TT-BNV, ngày 29/6/2021 của Bộ Nội vụ sửa đổi, bổ sung chế độ nâng bậc lương thường xuyên, nâng bậc lương trước thời hạn và chế độ phụ cấp thâm niên vượt khung đối với cán bộ, công chức, viên chức và người lao động;"
CAN_CU_PHAN_CAP = "Căn cứ Quy định số 09-QĐ/TU ngày 13/11/2025 của Tỉnh ủy về phân cấp quản lý cán bộ và quy hoạch, bổ nhiệm, giới thiệu ứng cử, tạm đình chỉ công tác, cho thôi giữ chức vụ, từ chức, miễn nhiệm đối với cán bộ;"
CAN_CU_CONG_VAN_BTC = "Căn cứ Công văn số 588-CV/BTCTU, ngày 10/12/2025 của Ban Tổ chức Tỉnh ủy về việc nâng bậc lương thường xuyên, nâng phụ cấp thâm niên vượt khung và nâng bậc lương trước thời hạn."

THANG_LA = {1:"một",2:"hai",3:"ba",4:"tư",5:"năm",6:"sáu",7:"bảy",8:"tám",9:"chín",10:"mười",11:"mười một",12:"mười hai"}


# ═════════════════════════════════════════════════════════
# 1. TIỆN ÍCH CHUNG
# ═════════════════════════════════════════════════════════
def _new_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)
    return doc


def _set_run(run, size=14, bold=False, italic=False, underline=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _p(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14, bold=False, italic=False,
       space_after=6, space_before=0, first_line_indent=None):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = 1.15
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(first_line_indent)
    if text:
        _set_run(para.add_run(text), size=size, bold=bold, italic=italic)
    return para


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)


def _quoc_hieu(doc, dong1_trai, dong2_trai_bold, ngay_ky, thang_ky, nam_ky, so_hieu=None, dong3_trai=None):
    """Bảng 2 cột: [Tên cơ quan ban hành / Số hiệu]  |  [Đảng Cộng sản VN / ngày tháng năm]"""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(7.5)
    table.columns[1].width = Cm(8.8)
    _remove_table_borders(table)

    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.width = Cm(7.5)
    right.width = Cm(8.8)

    def _cell_p(cell, text, bold=False, italic=False, first=False):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        _set_run(p.add_run(text), size=14, bold=bold, italic=italic)
        return p

    _cell_p(left, dong1_trai, first=True)
    _cell_p(left, dong2_trai_bold, bold=True)
    if dong3_trai:
        _cell_p(left, dong3_trai)
    _cell_p(left, "*")
    if so_hieu:
        _cell_p(left, so_hieu)

    _cell_p(right, "ĐẢNG CỘNG SẢN VIỆT NAM", bold=True, first=True)
    _cell_p(right, f"Tuyên Quang, ngày {ngay_ky} tháng {thang_ky} năm {nam_ky}", italic=True)
    return table


def _tieu_de_van_ban(doc, dong1, dong2=None):
    _p(doc, dong1, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=0)
    if dong2:
        _p(doc, dong2, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, space_after=0)
    _p(doc, "-----", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=False, space_after=10)


def _khoi_ky_ten(doc, noi_nhan_list, chuc_danh_ky, ten_nguoi_ky):
    """Bảng 2 cột không viền: Nơi nhận (trái) | Chức danh + Họ tên người ký (phải)"""
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(7.5)
    table.columns[1].width = Cm(8.8)
    _remove_table_borders(table)
    left = table.cell(0, 0)
    right = table.cell(0, 1)

    p0 = left.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    run0 = p0.add_run("Nơi nhận:")
    _set_run(run0, size=13, italic=False, underline=True)
    for item in noi_nhan_list:
        p = left.add_paragraph(f"- {item}")
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            _set_run(r, size=13)

    pr0 = right.paragraphs[0]
    pr0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pr0.add_run(chuc_danh_ky), size=14, bold=True)
    for _ in range(3):
        blank = right.add_paragraph()
        blank.alignment = WD_ALIGN_PARAGRAPH.CENTER
        blank.paragraph_format.space_after = Pt(0)
    pname = right.add_paragraph()
    pname.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pname.add_run(ten_nguoi_ky), size=14, bold=True)
    return table


def _save(doc):
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ═════════════════════════════════════════════════════════
# 2. XÁC ĐỊNH LOẠI NÂNG LƯƠNG / LÃNH ĐẠO TỪ DỮ LIỆU HÀNG (row)
#    — dùng trực tiếp các cột đã có trong df_calculated của app
# ═════════════════════════════════════════════════════════
def xac_dinh_loai_nang_luong(row):
    """Trả về 'vuot_khung' hoặc 'thuong_xuyen' dựa trên dữ liệu đã tính toán."""
    vk_ht = str(row.get('vuot_khung_hien_tai', '') or '').strip()
    if vk_ht.lower() not in ('', 'none', 'nan'):
        return "vuot_khung"
    vk_moi = str(row.get('vuot_khung_moi', '') or '').strip()
    bac_ht = str(row.get('bac_luong', '') or '').strip()
    bac_moi = str(row.get('bac_luong_moi', '') or '').strip()
    if vk_moi not in ('', 'none', 'nan', 'None') and bac_moi == bac_ht:
        return "vuot_khung"  # vượt khung lần đầu
    return "thuong_xuyen"


def la_lanh_dao(row):
    cv = str(row.get('chuc_vu', '') or '').upper()
    return ("TRƯỞNG BAN" in cv) or ("PHÓ TRƯỞNG BAN" in cv) or ("TRƯỞNG BAN" in cv.replace("PHÓ ", ""))


def _truong_dien_bien(row):
    """Trích xuất & định dạng các trường dữ liệu dùng chung để mô tả diễn biến lương."""
    ngach = str(row.get('ngach_luong', '') or '').strip()
    ma_ngach = str(row.get('ma_ngach', '') or '').strip()
    bac_ht = str(row.get('bac_luong', '') or '').strip()
    hs_ht = str(row.get('he_so_hien_tai', '') or '').strip()
    ngay_ht = str(row.get('ngay_gan_nhat', '') or '').strip()
    bac_moi = str(row.get('bac_luong_moi', '') or '').strip()
    hs_moi = str(row.get('he_so_moi', '') or '').strip()
    ngay_dk = str(row.get('ngay_du_kien', '') or '').strip()
    vk_ht = str(row.get('vuot_khung_hien_tai', '') or '').strip()
    vk_moi = str(row.get('vuot_khung_moi', '') or '').strip()
    try:
        thang_nam_ht = datetime.strptime(ngay_ht, '%d/%m/%Y').strftime('%m/%Y')
    except Exception:
        thang_nam_ht = ngay_ht
    try:
        thang_nam_moi = datetime.strptime(ngay_dk, '%d/%m/%Y').strftime('%m/%Y')
    except Exception:
        thang_nam_moi = ngay_dk
    return dict(ngach=ngach, ma_ngach=ma_ngach, bac_ht=bac_ht, hs_ht=hs_ht, ngay_ht=ngay_ht,
                bac_moi=bac_moi, hs_moi=hs_moi, ngay_dk=ngay_dk, vk_ht=vk_ht, vk_moi=vk_moi,
                thang_nam_ht=thang_nam_ht, thang_nam_moi=thang_nam_moi)


def cau_dien_bien_luong(row, loai):
    """Sinh câu mô tả diễn biến lương/phụ cấp cho 1 người (1 câu gộp) — dùng cho QĐ/TT và Biên bản nhiều người."""
    f = _truong_dien_bien(row)
    if loai == "vuot_khung":
        if f['vk_ht'].lower() not in ('', 'none', 'nan'):
            return (f"Lương đang hưởng ngạch {f['ngach']} (Mã số {f['ma_ngach']}), bậc {f['bac_ht']}, hệ số {f['hs_ht']} "
                    f"và hưởng phụ cấp thâm niên vượt khung {f['vk_ht']}, kể từ ngày {f['ngay_ht']}. "
                    f"Nay nâng mức phụ cấp thâm niên vượt khung lên {f['vk_moi']}, kể từ ngày {f['ngay_dk']}.")
        else:
            return (f"Lương đang hưởng ngạch {f['ngach']} (Mã số {f['ma_ngach']}) bậc {f['bac_ht']}; hệ số {f['hs_ht']} "
                    f"từ tháng {f['thang_nam_ht']}. Nay nâng phụ cấp thâm niên vượt khung lần đầu bằng {f['vk_moi']}, "
                    f"kể từ ngày {f['ngay_dk']}.")
    else:
        return (f"Lương đang hưởng ngạch {f['ngach']} (Mã số {f['ma_ngach']}), bậc {f['bac_ht']}, hệ số {f['hs_ht']} "
                f"từ tháng {f['thang_nam_ht']}. Nâng lên bậc {f['bac_moi']}, hệ số {f['hs_moi']} từ tháng {f['thang_nam_moi']}.")


def _dien_bien_2_doan(row, loai):
    """Như cau_dien_bien_luong nhưng tách thành 2 đoạn riêng (hiện trạng / nay nâng) —
    dùng cho Biên bản khi chỉ có 1 người, đúng bố cục 2 đoạn của mẫu gốc."""
    f = _truong_dien_bien(row)
    if loai == "vuot_khung":
        if f['vk_ht'].lower() not in ('', 'none', 'nan'):
            doan1 = (f"Lương đang hưởng ngạch {f['ngach']} (Mã số {f['ma_ngach']}), bậc {f['bac_ht']} hệ số {f['hs_ht']} "
                     f"và hưởng phụ cấp thâm niên vượt khung {f['vk_ht']}, kể từ ngày {f['ngay_ht']}.")
            doan2 = f"Nay nâng mức phụ cấp thâm niên vượt khung lên {f['vk_moi']}, kể từ ngày {f['ngay_dk']}."
        else:
            doan1 = (f"Lương đang hưởng ngạch {f['ngach']} (Mã số {f['ma_ngach']}) bậc {f['bac_ht']}; hệ số {f['hs_ht']} "
                     f"từ tháng {f['thang_nam_ht']}.")
            doan2 = f"Nay nâng phụ cấp thâm niên vượt khung lần đầu bằng {f['vk_moi']}, kể từ ngày {f['ngay_dk']}."
    else:
        doan1 = (f"Lương đang hưởng ngạch {f['ngach']} (Mã số {f['ma_ngach']}), bậc {f['bac_ht']}, hệ số {f['hs_ht']} "
                 f"từ tháng {f['thang_nam_ht']}.")
        doan2 = f"Nâng lên bậc {f['bac_moi']}, hệ số {f['hs_moi']} từ tháng {f['thang_nam_moi']}."
    return doan1, doan2


# ═════════════════════════════════════════════════════════
# 3. QUYẾT ĐỊNH NÂNG LƯƠNG (cho công chức không phải Trưởng/Phó Trưởng ban)
# ═════════════════════════════════════════════════════════
def tao_quyet_dinh(r, loai, so_qd, ngay_ky, thang_ky, nam_ky, ngay_hop_bb,
                    truong_ban="Trần Mạnh Lợi", phong_lien_quan="Văn phòng Ban"):
    """
    r: dict dữ liệu MỘT cán bộ (đã tinh_toan_nang_luong). Mỗi công chức có 1 Quyết định riêng
       (số hiệu, hiệu lực pháp lý độc lập) — hàm này không gộp nhiều người vào 1 văn bản.
    loai: 'thuong_xuyen' | 'vuot_khung'
    """
    doc = _new_doc()
    tieu_de_1 = "nâng lương thường xuyên đối với công chức" if loai == "thuong_xuyen" \
        else "nâng phụ cấp thâm niên vượt khung đối với công chức"

    _quoc_hieu(doc, "TỈNH UỶ TUYÊN QUANG", "BAN TUYÊN GIÁO VÀ DÂN VẬN", ngay_ky, thang_ky, nam_ky,
               so_hieu=f"Số {so_qd}-QĐ/BTGDVTU")
    doc.add_paragraph()
    _tieu_de_van_ban(doc, "QUYẾT ĐỊNH", tieu_de_1)

    can_cu = CAN_CU_THUONG_XUYEN if loai == "thuong_xuyen" else CAN_CU_VUOT_KHUNG
    _p(doc, can_cu)
    _p(doc, CAN_CU_PHAN_CAP)
    _p(doc, f"Căn cứ Biên bản cuộc họp ngày {ngay_hop_bb} của Hội đồng xét nâng bậc lương cơ quan "
            f"Ban Tuyên giáo và Dân vận Tỉnh ủy;")
    _p(doc, "Xét đề nghị của Chánh Văn phòng Ban.")
    doc.add_paragraph()

    _p(doc, "BAN TUYÊN GIÁO VÀ DÂN VẬN TỈNH UỶ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _p(doc, "QUYẾT ĐỊNH", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=10)

    ho_ten = str(r.get('ho_ten', ''))
    chuc_vu = str(r.get('chuc_vu', '') or '').strip()
    para = _p(doc, "")
    run1 = para.add_run("Điều 1. ")
    _set_run(run1, bold=True)
    run2 = para.add_run(f"Đồng chí {ho_ten}" + (f", {chuc_vu}" if chuc_vu else "") + ". "
                         + cau_dien_bien_luong(r, loai))
    _set_run(run2)
    _p(doc, f"Thời gian nâng bậc lương lần sau tính từ ngày {r.get('ngay_du_kien','')}.")
    nguoi_thi_hanh = f"Phòng {chuc_vu.split('Phòng')[-1].strip()}" if 'Phòng' in chuc_vu else phong_lien_quan
    cau_dieu2 = (f"{phong_lien_quan}, {nguoi_thi_hanh} và đồng chí {ho_ten} căn cứ Quyết định thi hành."
                 if nguoi_thi_hanh != phong_lien_quan else
                 f"{phong_lien_quan} và đồng chí {ho_ten} căn cứ Quyết định thi hành.")

    para2 = _p(doc, "")
    run1 = para2.add_run("Điều 2. ")
    _set_run(run1, bold=True)
    run2 = para2.add_run(cau_dieu2)
    _set_run(run2)
    doc.add_paragraph()

    noi_nhan = ["Như điều 2,", "Kế toán Ban,", "Hồ sơ cán bộ,", "Lưu Ban Tuyên giáo và Dân vận Tỉnh ủy."]
    _khoi_ky_ten(doc, noi_nhan, "TRƯỞNG BAN", truong_ban)
    return _save(doc)


# ═════════════════════════════════════════════════════════
# 4. TỜ TRÌNH GỬI BAN TỔ CHỨC TỈNH ỦY (cho Trưởng/Phó Trưởng ban)
# ═════════════════════════════════════════════════════════
def tao_to_trinh(ds_lanh_dao, loai, so_tt, ngay_ky, thang_ky, nam_ky, ngay_hop_bb,
                  truong_ban="Trần Mạnh Lợi"):
    doc = _new_doc()
    tieu_de_1 = "V/v đề nghị nâng bậc lương thường xuyên" if loai == "thuong_xuyen" \
        else "V/v đề nghị nâng phụ cấp thâm niên vượt khung"

    _quoc_hieu(doc, "TỈNH ỦY TUYÊN QUANG", "BAN TUYÊN GIÁO VÀ DÂN VẬN", ngay_ky, thang_ky, nam_ky,
               so_hieu=f"Số {so_tt}-TTr/BTGDV")
    doc.add_paragraph()
    _tieu_de_van_ban(doc, "TỜ TRÌNH", tieu_de_1)

    kg = doc.add_paragraph()
    kg.paragraph_format.space_after = Pt(8)
    run_kg = kg.add_run("Kính gửi: ")
    _set_run(run_kg, italic=True)
    run_kg2 = kg.add_run("Ban Tổ chức Tỉnh ủy.")
    _set_run(run_kg2, bold=True)

    can_cu = CAN_CU_VUOT_KHUNG if loai == "vuot_khung" else CAN_CU_THUONG_XUYEN
    doan1 = f"{can_cu} {CAN_CU_PHAN_CAP} {CAN_CU_CONG_VAN_BTC}"
    _p(doc, doan1)
    _p(doc, f"Căn cứ kết quả đánh giá xếp loại cán bộ, công chức năm {int(nam_ky)-1}; Biên bản cuộc họp Hội đồng "
            f"xét nâng lương của Ban Tuyên giáo và Dân vận Tỉnh ủy ngày {ngay_hop_bb};")

    nhan_dong = "nâng bậc lương thường xuyên" if loai == "thuong_xuyen" else "nâng phụ cấp thâm niên vượt khung"
    if len(ds_lanh_dao) == 1:
        r = ds_lanh_dao[0]
        ho_ten = str(r.get('ho_ten', ''))
        chuc_vu = str(r.get('chuc_vu', '') or '').strip()
        _p(doc, f"Ban Tuyên giáo và Dân vận Tỉnh ủy đề nghị Ban Tổ chức Tỉnh ủy thẩm định, trình Thường trực "
                f"Tỉnh ủy xem xét Quyết định {nhan_dong} cho đồng chí {ho_ten}"
                + (f", {chuc_vu}" if chuc_vu else "") + ".")
        _p(doc, cau_dien_bien_luong(r, loai))
    else:
        _p(doc, f"Ban Tuyên giáo và Dân vận Tỉnh ủy đề nghị Ban Tổ chức Tỉnh ủy thẩm định, trình Thường trực "
                f"Tỉnh ủy xem xét, Quyết định {nhan_dong} cho các đồng chí có tên sau:")
        for i, r in enumerate(ds_lanh_dao, 1):
            ho_ten = str(r.get('ho_ten', ''))
            chuc_vu = str(r.get('chuc_vu', '') or '').strip()
            p = _p(doc, "", space_after=6)
            run = p.add_run(f"{i}. Đồng chí {ho_ten}" + (f", {chuc_vu}" if chuc_vu else "") + ". "
                             + cau_dien_bien_luong(r, loai))
            _set_run(run)

    ghi_chu = _p(doc, "(Có Biên bản họp xét và dự thảo Quyết định kèm theo)",
                 align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=8)
    _p(doc, "Ban Tuyên giáo và Dân vận Tỉnh ủy trân trọng đề nghị Ban Tổ chức Tỉnh ủy trình Thường trực Tỉnh ủy "
            "xem xét, quyết định.")
    doc.add_paragraph()

    noi_nhan = ["Như kính gửi,", "Lãnh đạo Ban,", "Lưu Ban Tuyên giáo và Dân vận Tỉnh ủy."]
    _khoi_ky_ten(doc, noi_nhan, "TRƯỞNG BAN", truong_ban)
    return _save(doc)


# ═════════════════════════════════════════════════════════
# 5. BIÊN BẢN HỌP HỘI ĐỒNG XÉT NÂNG LƯƠNG
#    — mỗi biên bản CHỈ gồm 1 loại nâng lương (thường xuyên HOẶC vượt khung)
#      và 1 nhóm đối tượng (lãnh đạo Ban trình BTC Tỉnh ủy HOẶC cán bộ thường
#      do Trưởng Ban trực tiếp quyết định) — không trộn lẫn.
# ═════════════════════════════════════════════════════════
def tao_bien_ban(ds_nhom, loai, la_lanh_dao_nhom, ngay_hop, gio_bat_dau, gio_ket_thuc, thanh_phan_text,
                  truong_ban="Trần Mạnh Lợi", thu_ky="Đinh Thị Thúy",
                  dia_diem="Phòng họp Hội đồng xét nâng lương Ban Tuyên giáo và Dân vận Tỉnh ủy"):
    """
    ds_nhom: list[dict] CHỈ các cán bộ CÙNG loại `loai` và CÙNG thuộc nhóm `la_lanh_dao_nhom`.
    loai: 'thuong_xuyen' | 'vuot_khung'
    la_lanh_dao_nhom: True nếu nhóm là Trưởng/Phó Trưởng ban (đề nghị BTC Tỉnh ủy),
                       False nếu là cán bộ, chuyên viên khác (Trưởng Ban trực tiếp quyết định).
    """
    doc = _new_doc()

    # Sắp xếp: lãnh đạo Ban lên trước, kế đến lãnh đạo phòng, sau cùng là chuyên viên/nhân viên khác.
    def _thu_tu(row):
        cv = str(row.get('chuc_vu', '') or '').upper()
        if la_lanh_dao(row):
            return 0
        if 'TRƯỞNG PHÒNG' in cv:
            return 1
        return 2
    ds_nhom = sorted(ds_nhom, key=_thu_tu)

    hanh_dong = "nâng bậc lương thường xuyên" if loai == "thuong_xuyen" else "nâng phụ cấp thâm niên vượt khung"
    tieu_de_2 = f"Họp xét {hanh_dong} cho cán bộ, công chức"
    can_cu = CAN_CU_THUONG_XUYEN if loai == "thuong_xuyen" else CAN_CU_VUOT_KHUNG
    can_cu_ngan = can_cu[len("Căn cứ "):] if can_cu.startswith("Căn cứ ") else can_cu

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(7.5)
    table.columns[1].width = Cm(8.8)
    _remove_table_borders(table)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    p0 = left.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p0.add_run("TỈNH ỦY TUYÊN QUANG"), size=14)
    p1 = left.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p1.add_run("BAN TUYÊN GIÁO VÀ DÂN VẬN"), size=14, bold=True)
    p2 = left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p2.add_run("*"), size=14)
    pr0 = right.paragraphs[0]
    pr0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pr0.add_run("ĐẢNG CỘNG SẢN VIỆT NAM"), size=14, bold=True)
    pr1 = right.add_paragraph()
    pr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        ngay_dt = datetime.strptime(ngay_hop, '%d/%m/%Y')
        ngay_s, thang_s, nam_s = ngay_dt.strftime('%d'), ngay_dt.strftime('%m'), ngay_dt.strftime('%Y')
    except Exception:
        ngay_s, thang_s, nam_s = ngay_hop, "", ""
    _set_run(pr1.add_run(f"Tuyên Quang, ngày {ngay_s} tháng {thang_s} năm {nam_s}"), size=14, italic=True)
    nam_truoc = str(int(nam_s) - 1) if nam_s else ""

    doc.add_paragraph()
    _tieu_de_van_ban(doc, "BIÊN BẢN", tieu_de_2)

    # ── Đoạn mở đầu ─────────────────────────────
    if len(ds_nhom) == 1:
        r0 = ds_nhom[0]
        ho_ten0 = str(r0.get('ho_ten', ''))
        chuc_vu0 = str(r0.get('chuc_vu', '') or '').strip()
        _p(doc, f"Vào hồi {gio_bat_dau} phút, ngày {ngay_hop} tại {dia_diem} họp xét {hanh_dong} cho "
                f"đồng chí {ho_ten0}" + (f", {chuc_vu0}." if chuc_vu0 else "."))
    else:
        ten_ds = "; ".join(
            f"đồng chí {r.get('ho_ten','')}" + (f", {r.get('chuc_vu','').strip()}" if str(r.get('chuc_vu','')).strip() else "")
            for r in ds_nhom
        )
        _p(doc, f"Vào hồi {gio_bat_dau} phút, ngày {ngay_hop} tại {dia_diem} họp xét {hanh_dong} cho "
                f"{len(ds_nhom):02d} công chức, gồm: {ten_ds}.")

    _p(doc, "I - THÀNH PHẦN", bold=True, space_after=6)
    for line in [l.strip() for l in thanh_phan_text.split("\n") if l.strip()]:
        prefix = "" if line.startswith("-") or line.startswith("Đồng chí") or line.startswith("Đ/c") else "- "
        _p(doc, f"{prefix}{line}" if not line.startswith("-") else line, space_after=4)

    # ── Nội dung ─────────────────────────────
    _p(doc, "II- NỘI DUNG", bold=True, space_after=6)
    p1 = _p(doc, "")
    r1 = p1.add_run("1. ")
    _set_run(r1, bold=True)
    r1b = p1.add_run(f"Đồng chí chủ trì thông qua {can_cu_ngan}")
    _set_run(r1b)

    doi_tuong_txt = (f"của đồng chí {ds_nhom[0].get('ho_ten','')}" if len(ds_nhom) == 1
                      else "của các đồng chí có tên trên")
    p2 = _p(doc, "")
    r2 = p2.add_run("2. ")
    _set_run(r2, bold=True)
    r2b = p2.add_run(f"Hội đồng tiến hành rà soát tiêu chuẩn, điều kiện {hanh_dong}, đồng thời đối chiếu với "
                      f"kết quả nhận xét, đánh giá cán bộ năm {nam_truoc} {doi_tuong_txt}.")
    _set_run(r2b)

    ket_luan_dau = ("Hội đồng xét nâng lương Ban Tuyên giáo và Dân vận Tỉnh ủy biểu quyết thống nhất đề nghị "
                     "Ban Tổ chức Tỉnh ủy thẩm định, trình Thường trực Tỉnh ủy xem xét, Quyết định"
                     if la_lanh_dao_nhom else
                     "Hội đồng xét nâng lương Ban Tuyên giáo và Dân vận Tỉnh ủy biểu quyết thống nhất đề nghị "
                     "Trưởng Ban Quyết định")

    if len(ds_nhom) == 1:
        r0 = ds_nhom[0]
        ho_ten0 = str(r0.get('ho_ten', ''))
        chuc_vu0 = str(r0.get('chuc_vu', '') or '').strip()
        loai0 = r0.get('loai') or loai
        _p(doc, f"Sau khi thảo luận, xem xét, {ket_luan_dau} {hanh_dong} cho đồng chí {ho_ten0}"
                + (f", {chuc_vu0}" if chuc_vu0 else "") + ", cụ thể như sau:")
        doan1, doan2 = _dien_bien_2_doan(r0, loai0)
        _p(doc, doan1)
        _p(doc, doan2)
    else:
        _p(doc, f"Sau khi thảo luận, xem xét, {ket_luan_dau} {hanh_dong} cho các đồng chí có tên sau, "
                f"cụ thể như sau:")
        for i, r in enumerate(ds_nhom, 1):
            loai_r = r.get('loai') or loai
            ho_ten = str(r.get('ho_ten', ''))
            chuc_vu = str(r.get('chuc_vu', '') or '').strip()
            p = _p(doc, "", space_after=6)
            rn = p.add_run(f"{i}. ")
            _set_run(rn, bold=True)
            rb = p.add_run(f"Đồng chí {ho_ten}" + (f", {chuc_vu}" if chuc_vu else "") + ". "
                            + cau_dien_bien_luong(r, loai_r))
            _set_run(rb)

    _p(doc, f"Cuộc họp kết thúc vào hồi {gio_ket_thuc} phút cùng ngày. Biên bản đã được thông qua trước Hội đồng "
            "và được 100% thành viên nhất trí.")
    doc.add_paragraph()

    table2 = doc.add_table(rows=1, cols=2)
    table2.columns[0].width = Cm(7.5)
    table2.columns[1].width = Cm(8.8)
    _remove_table_borders(table2)
    l2 = table2.cell(0, 0)
    r2c = table2.cell(0, 1)
    pl = l2.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pl.add_run("THƯ KÝ"), size=14, bold=True)
    for _ in range(3):
        b = l2.add_paragraph()
        b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pln = l2.add_paragraph()
    pln.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pln.add_run(thu_ky), size=14, bold=True)

    pr = r2c.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(pr.add_run("CHỦ TRÌ"), size=14, bold=True)
    for _ in range(3):
        b = r2c.add_paragraph()
        b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prn = r2c.add_paragraph()
    prn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(prn.add_run(truong_ban), size=14, bold=True)

    return _save(doc)
