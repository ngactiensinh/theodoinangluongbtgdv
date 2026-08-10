import streamlit as st
import pandas as pd
from supabase import create_client
from datetime import datetime
from dateutil.relativedelta import relativedelta
import json
import io
import plotly.express as px
import plotly.graph_objects as go

from van_ban_nang_luong import (
    tao_quyet_dinh, tao_to_trinh, tao_bien_ban,
    xac_dinh_loai_nang_luong, la_lanh_dao, THANH_PHAN_HOI_DONG_MAC_DINH,
)

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except:
    pass

try:
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    try:
        from openpyxl.utils import get_column_letter
    except ImportError:
        from openpyxl.utils.cell import get_column_letter
except:
    pass

# ─────────────────────────────────────────
# 1. CẤU HÌNH TRANG & CSS
# ─────────────────────────────────────────
st.set_page_config(
    page_title="Quản lý Lương – Ban TG&DV Tuyên Quang",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Be+Vietnam+Pro:wght@300;400;500;600;700;800&display=swap');
:root {
    --primary:   #0d47a1; --primary-l: #1565c0; --accent:    #c62828;
    --teal:      #00838f; --bg:        #f4f6fb; --card:      #ffffff;
    --border:    #dde3f0; --text:      #1a2035; --muted:     #6b7a99;
    --radius:    12px;    --shadow:    0 2px 12px rgba(13,71,161,.08);
}
html, body, [class*="css"] { font-family: 'Be Vietnam Pro', sans-serif !important; }
.main .block-container { padding: 1.2rem 2rem 2rem; max-width: 1400px; }
.stApp { background: var(--bg); }
[data-testid="stSidebar"] { background: linear-gradient(180deg, #0d1b3e 0%, #1a2f6e 60%, #0d47a1 100%) !important; border-right: none !important; }
[data-testid="stSidebar"] * { color: #e8eaf6 !important; }
[data-testid="stSidebar"] .stTextInput input { background: rgba(255,255,255,.1) !important; border: 1px solid rgba(255,255,255,.25) !important; color: #fff !important; border-radius: 8px !important; }
[data-testid="stSidebar"] .stButton button { background: rgba(255,255,255,.12) !important; border: 1px solid rgba(255,255,255,.3) !important; color: #fff !important; border-radius: 8px !important; font-weight: 600 !important; transition: all .2s; }
[data-testid="stSidebar"] .stButton button:hover { background: rgba(255,255,255,.22) !important; }
.hero { background: linear-gradient(135deg, #0d47a1 0%, #1565c0 50%, #00838f 100%); padding: 22px 28px; border-radius: var(--radius); color: #fff; display: flex; align-items: center; gap: 20px; margin-bottom: 20px; box-shadow: 0 4px 24px rgba(13,71,161,.25); position: relative; overflow: hidden; }
.hero::before { content: ''; position: absolute; inset: 0; background: repeating-linear-gradient(45deg, transparent, transparent 40px, rgba(255,255,255,.03) 40px, rgba(255,255,255,.03) 80px); }
.hero-icon { font-size: 42px; line-height: 1; flex-shrink: 0; }
.hero-title { font-size: 22px; font-weight: 800; letter-spacing: -.3px; margin: 0; }
.hero-sub   { font-size: 13px; opacity: .85; margin: 3px 0 0; }
.hero-badge { margin-left: auto; background: rgba(255,255,255,.15); backdrop-filter: blur(6px); border: 1px solid rgba(255,255,255,.25); border-radius: 20px; padding: 6px 14px; font-size: 12px; font-weight: 600; white-space: nowrap; }
.kpi-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 14px; margin-bottom: 20px; }
.kpi { background: var(--card); border: 1px solid var(--border); border-radius: var(--radius); padding: 18px 20px; box-shadow: var(--shadow); position: relative; overflow: hidden; }
.kpi::after { content: ''; position: absolute; top: 0; left: 0; right: 0; height: 3px; border-radius: var(--radius) var(--radius) 0 0; }
.kpi.blue::after  { background: var(--primary); }
.kpi.red::after   { background: var(--accent); }
.kpi.teal::after  { background: var(--teal); }
.kpi.gold::after  { background: #f57f17; }
.kpi-label { font-size: 11px; font-weight: 700; color: var(--muted); text-transform: uppercase; letter-spacing: .6px; margin-bottom: 6px; }
.kpi-val   { font-size: 36px; font-weight: 800; line-height: 1; color: var(--primary); }
.kpi.red  .kpi-val { color: var(--accent); }
.kpi.teal .kpi-val { color: var(--teal); }
.kpi.gold .kpi-val { color: #f57f17; }
.kpi-icon { position: absolute; right: 16px; top: 14px; font-size: 28px; opacity: .15; }
[data-testid="stDataFrame"], [data-testid="data-editor"] { border: 1px solid var(--border) !important; border-radius: var(--radius) !important; overflow: hidden !important; box-shadow: var(--shadow) !important; }
.stTabs [data-baseweb="tab-list"] { background: var(--card) !important; border-radius: var(--radius) var(--radius) 0 0 !important; border: 1px solid var(--border) !important; border-bottom: none !important; padding: 6px 8px 0 !important; gap: 4px !important; }
.stTabs [data-baseweb="tab"] { background: transparent !important; border-radius: 8px 8px 0 0 !important; padding: 8px 18px !important; font-weight: 600 !important; font-size: 14px !important; color: var(--muted) !important; border: none !important; }
.stTabs [aria-selected="true"] { background: var(--primary) !important; color: #fff !important; }
.stTabs [data-baseweb="tab-panel"] { background: var(--card) !important; border: 1px solid var(--border) !important; border-top: none !important; border-radius: 0 0 var(--radius) var(--radius) !important; padding: 20px !important; }
.stDownloadButton button, .stButton button { border-radius: 8px !important; font-weight: 600 !important; font-size: 13px !important; padding: 8px 16px !important; transition: all .2s !important; }
.stDownloadButton button { background: var(--primary) !important; color: #fff !important; border: none !important; }
.stDownloadButton button:hover { background: var(--primary-l) !important; transform: translateY(-1px) !important; box-shadow: 0 4px 12px rgba(13,71,161,.3) !important; }
.section-title { font-size: 13px; font-weight: 700; color: var(--muted); text-transform: uppercase; letter-spacing: .8px; margin: 0 0 10px; padding-bottom: 8px; border-bottom: 2px solid var(--border); }
.admin-badge { background: linear-gradient(135deg, #1b5e20, #2e7d32); color: #fff !important; padding: 10px 14px; border-radius: 10px; font-size: 13px; font-weight: 700; text-align: center; margin-bottom: 14px; border: 1px solid rgba(255,255,255,.2); }
.user-badge { background: rgba(255,255,255,.08); border: 1px solid rgba(255,255,255,.2); padding: 10px 14px; border-radius: 10px; font-size: 12px; text-align: center; margin-bottom: 14px; color: rgba(255,255,255,.7) !important; }
.divider { border: none; border-top: 1px solid var(--border); margin: 16px 0; }
label { font-size: 12px !important; font-weight: 600 !important; color: var(--muted) !important; }
.chart-wrap { background: var(--card); border: 1px solid var(--border); border-radius: var(--radius); padding: 14px; box-shadow: var(--shadow); }
.vb-group { background: #f4f6fb; border: 1px solid var(--border); border-radius: 10px; padding: 12px 16px; margin-bottom: 10px; }
.vb-group b { color: var(--primary); }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────
# 2. SESSION STATE & PHÂN QUYỀN
# ─────────────────────────────────────────
if "role" not in st.session_state:
    st.session_state.role = "user"

with st.sidebar:
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align:center; margin-bottom:20px;">
        <div style="font-size:48px;">⭐</div>
        <div style="font-size:14px; font-weight:700; opacity:.9; line-height:1.4;">
            BAN TUYÊN GIÁO<br>& DÂN VẬN
        </div>
        <div style="font-size:11px; opacity:.6; margin-top:4px;">Tỉnh ủy Tuyên Quang</div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("<hr>", unsafe_allow_html=True)

    if st.session_state.role == "user":
        st.markdown('<div class="user-badge">👁 Chế độ: XEM (chỉ đọc)</div>', unsafe_allow_html=True)
        with st.expander("🔐 Đăng nhập Quản trị"):
            admin_pwd = st.text_input("Mật khẩu:", type="password", label_visibility="collapsed", placeholder="Nhập mật khẩu admin...")
            if st.button("Đăng nhập", use_container_width=True):
                if admin_pwd == "Admin@2026":
                    st.session_state.role = "admin"
                    st.success("✅ Đăng nhập thành công!")
                    st.rerun()
                else:
                    st.error("❌ Sai mật khẩu!")
    else:
        st.markdown('<div class="admin-badge">✅ QUYỀN ADMIN ĐÃ KÍCH HOẠT</div>', unsafe_allow_html=True)
        if st.button("🚪 Đăng xuất", use_container_width=True):
            st.session_state.role = "user"
            st.rerun()

    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown("""
    <div style="font-size:11px; opacity:.5; text-align:center; line-height:1.8;">
        Phiên bản 4.6<br>
        Phát triển bởi Tuấn 🚀<br>
        © 2025 Ban TG&DV Tuyên Quang
    </div>
    """, unsafe_allow_html=True)

# ─────────────────────────────────────────
# 4. HEADER BANNER
# ─────────────────────────────────────────
role_label = "🔴 ADMIN" if st.session_state.role == "admin" else "👁 XEM"
st.markdown(f"""
<div class="hero">
    <div class="hero-icon">📊</div>
    <div>
        <div class="hero-title">HỆ THỐNG QUẢN LÝ LƯƠNG 4.0</div>
        <div class="hero-sub">Ban Tuyên giáo và Dân vận – Tỉnh ủy Tuyên Quang</div>
    </div>
    <div class="hero-badge">{role_label}</div>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────
# 5. KẾT NỐI SUPABASE
# ─────────────────────────────────────────
try:
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    supabase = create_client(url, key)
except Exception as e:
    st.error("⚠️ Lỗi kết nối cơ sở dữ liệu. Vui lòng kiểm tra Secrets!")
    st.stop()

try:
    supabase.table("thong_ke_truy_cap").insert({"ten_app": "Theo dõi Nâng lương"}).execute()
except:
    pass

# ─────────────────────────────────────────
# 6. HÀM XỬ LÝ DỮ LIỆU
# ─────────────────────────────────────────
def format_ma_ngach(val):
    if pd.isna(val) or val == "" or str(val).lower() == "nan":
        return ""
    val_str = str(val).strip()
    return val_str[:-2] if val_str.endswith(".0") else val_str

def tinh_toan_nang_luong(df):
    res = df.copy()
    if res.empty:
        return res
    today = datetime.now().date()
    for idx, row in res.iterrows():
        ngach    = str(row.get('ngach_luong', '')).strip().upper()
        chuc_vu  = str(row.get('chuc_vu', '')).strip().upper()
        bac_ht   = str(row.get('bac_luong', '')).strip()
        hs_str   = str(row.get('he_so_hien_tai', '0')).replace(',', '.')
        try:
            hs_ht = float(hs_str)
        except:
            hs_ht = 0.0
        vk_ht        = str(row.get('vuot_khung_hien_tai', 'None')).strip()
        ngay_ht_str  = str(row.get('ngay_gan_nhat', ''))
        try:
            ngay_ht = datetime.strptime(ngay_ht_str, '%d/%m/%Y').date()
        except:
            for col in ['bac_luong_moi', 'he_so_moi', 'vuot_khung_moi', 'ngay_du_kien']:
                res.at[idx, col] = ""
            res.at[idx, 'trang_thai'] = "Chưa có ngày"
            continue

        is_vk  = (vk_ht.lower() != 'none' and '%' in vk_ht)
        bac_moi, hs_moi, vk_moi, ngay_dk = bac_ht, hs_ht, vk_ht, ngay_ht

        if is_vk:
            vk_val  = int(vk_ht.replace('%', '').strip())
            ngay_dk = ngay_ht + relativedelta(years=1)
            vk_moi  = f"{vk_val + 1}%"
        else:
            try:
                if '/' in bac_ht:
                    x, y = map(int, bac_ht.split('/'))
                else:
                    x, y = int(bac_ht), 99
                if x >= y:
                    ngay_dk = ngay_ht + relativedelta(years=3)
                    vk_moi  = "5%"
                else:
                    bac_moi  = f"{x+1}/{y}"

                    # Xác định nhóm ngạch để lấy đúng hệ số chênh lệch giữa 2 bậc liền kề
                    # và thời hạn nâng bậc lương thường xuyên tương ứng.
                    la_ke_toan_tc = 'KẾ TOÁN VIÊN TRUNG CẤP' in ngach or 'KẾ TOÁN VIÊN TRUNG CẤP' in chuc_vu
                    la_nhan_vien  = any(k in ngach or k in chuc_vu for k in ['LÁI XE', 'PHỤC VỤ', 'VĂN THƯ'])

                    if 'CVCC' in ngach:                 # Chuyên viên cao cấp — kiểm tra TRƯỚC 'CVC'
                        delta, interval = 0.62, 3
                    elif 'CVC' in ngach:                # Chuyên viên chính
                        delta, interval = 0.34, 3
                    elif la_ke_toan_tc:                 # Kế toán viên trung cấp (Loại B)
                        delta, interval = 0.20, 2
                    elif la_nhan_vien:                  # Lái xe / Phục vụ / Văn thư
                        delta, interval = 0.20, 2        # TODO: nhờ bạn xác nhận lại hệ số chính xác cho nhóm này
                    else:                                # Chuyên viên và tương đương (mặc định)
                        delta, interval = 0.33, 3

                    ngay_dk  = ngay_ht + relativedelta(years=interval)
                    hs_moi   = hs_ht + delta
            except:
                pass

        res.at[idx, 'bac_luong_moi']  = bac_moi
        res.at[idx, 'he_so_moi']      = f"{hs_moi:.2f}".replace('.', ',')
        res.at[idx, 'vuot_khung_moi'] = vk_moi
        res.at[idx, 'ngay_du_kien']   = ngay_dk.strftime('%d/%m/%Y')

        days_left = (ngay_dk - today).days
        if days_left < 0:
            res.at[idx, 'trang_thai'] = "⛔ Đã quá hạn"
        elif days_left <= 30:
            res.at[idx, 'trang_thai'] = "🔴 Sắp đến hạn (Tháng này)"
        elif days_left <= 90:
            res.at[idx, 'trang_thai'] = "🟡 Sắp đến hạn (Quý này)"
        else:
            res.at[idx, 'trang_thai'] = "🟢 Chưa đến hạn"
    return res.fillna("")

def style_trang_thai(val):
    v = str(val)
    if "quá hạn" in v.lower():  return 'color:#b71c1c; font-weight:700;'
    if "tháng này" in v.lower(): return 'color:#e65100; font-weight:700;'
    if "quý này" in v.lower():   return 'color:#f57f17; font-weight:600;'
    return 'color:#2e7d32;'

def tao_file_word_dien_bien(df, thang_chon="Tất cả", nam_chon="Tất cả"):
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for m in ['left', 'right', 'top', 'bottom']:
        setattr(section, f'{m}_margin', Cm(1.5))

    table_h = doc.add_table(rows=1, cols=2)
    table_h.columns[0].width = Cm(10)
    table_h.columns[1].width = Cm(16)
    cl = table_h.cell(0, 0).paragraphs[0]
    cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cl.add_run("TỈNH UỶ TUYÊN QUANG\n").bold = True
    cl.add_run("BAN TUYÊN GIÁO VÀ DÂN VẬN\n").bold = True
    cl.add_run("*")
    cr = table_h.cell(0, 1).paragraphs[0]
    cr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr.add_run("ĐẢNG CỘNG SẢN VIỆT NAM\n").bold = True

    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_t.add_run("\nBIỂU TỔNG HỢP DIỄN BIẾN LƯƠNG\n")
    run_t.bold = True
    run_t.font.size = Pt(14)
    txt_thang = thang_chon if thang_chon != "Tất cả" else datetime.now().strftime('%m')
    txt_nam   = nam_chon   if nam_chon   != "Tất cả" else datetime.now().strftime('%Y')
    p_t.add_run(f"Ban Tuyên giáo và Dân vận Tỉnh uỷ tháng {txt_thang} năm {txt_nam}").italic = True

    table = doc.add_table(rows=1, cols=11)
    table.style = 'Table Grid'
    headers = ['TT', 'Họ và tên', 'Chức vụ', 'Mã ngạch', 'Bậc HT',
               'Hệ số HT', 'Ngày hưởng', 'Bậc mới', 'Hệ số mới', 'Hưởng từ', 'Ghi chú']
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc_pr = c._tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '0D47A1')
            tc_pr.append(shd)
        except:
            pass

    for idx, (_, r) in enumerate(df.iterrows(), 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = str(r.get('ho_ten', ''))
        row_cells[2].text = str(r.get('chuc_vu', ''))
        row_cells[3].text = str(r.get('ma_ngach', ''))
        row_cells[4].text = str(r.get('bac_luong', ''))
        hs_ht = str(r.get('he_so_hien_tai', ''))
        if r.get('vuot_khung_hien_tai', '') not in ['', 'None', 'nan']:
            hs_ht += f" (VK {r.get('vuot_khung_hien_tai')})"
        row_cells[5].text = hs_ht
        row_cells[6].text = str(r.get('ngay_gan_nhat', ''))
        row_cells[7].text = str(r.get('bac_luong_moi', ''))
        hs_m = str(r.get('he_so_moi', ''))
        if r.get('vuot_khung_moi', '') not in ['', 'None', 'nan']:
            hs_m += f" (VK {r.get('vuot_khung_moi')})"
        row_cells[8].text = hs_m
        row_cells[9].text  = str(r.get('ngay_du_kien', ''))
        row_cells[10].text = "Nâng lương TX"
        for i in range(11):
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ─────────────────────────────────────────
# 7. MAIN
# ─────────────────────────────────────────
def main():
    try:
        res          = supabase.table("theo_doi_luong").select("*").execute()
        df_base      = pd.DataFrame(res.data) if res.data else pd.DataFrame(
            columns=["ho_ten", "ngay_gan_nhat", "ma_ngach"])
        
        df_calculated = tinh_toan_nang_luong(df_base)
        df_calculated = df_calculated[df_calculated['ho_ten'].astype(str).str.strip() != ""]

        # ── KPI CARDS ──────────────────────────────
        tong_nv     = len(df_calculated)
        qua_han     = len(df_calculated[df_calculated['trang_thai'] == "⛔ Đã quá hạn"])
        thang_nay   = len(df_calculated[df_calculated['trang_thai'] == "🔴 Sắp đến hạn (Tháng này)"])
        quy_nay     = len(df_calculated[df_calculated['trang_thai'] == "🟡 Sắp đến hạn (Quý này)"])

        st.markdown(f"""
        <div class="kpi-grid">
            <div class="kpi blue">
                <div class="kpi-icon">👥</div>
                <div class="kpi-label">Tổng cán bộ</div>
                <div class="kpi-val">{tong_nv}</div>
            </div>
            <div class="kpi red">
                <div class="kpi-icon">⛔</div>
                <div class="kpi-label">Đã quá hạn</div>
                <div class="kpi-val">{qua_han}</div>
            </div>
            <div class="kpi gold" style="--accent-color:#f57f17">
                <div class="kpi-icon">🔴</div>
                <div class="kpi-label">Tháng này</div>
                <div class="kpi-val" style="color:#e65100">{thang_nay}</div>
            </div>
            <div class="kpi teal">
                <div class="kpi-icon">🟡</div>
                <div class="kpi-label">Quý này</div>
                <div class="kpi-val">{quy_nay}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # ── TABS ───────────────────────────────────
        if st.session_state.role == "admin":
            tab1, tab2, tab3 = st.tabs(["📋  Quản lý & Lọc dữ liệu", "📊  Dashboard thống kê", "🏛️  Văn bản Nâng lương"])
        else:
            tab1, tab2 = st.tabs(["📋  Quản lý & Lọc dữ liệu", "📊  Dashboard thống kê"])

        # ══════════════════════════════════════════
        # TAB 1 – QUẢN LÝ
        # ══════════════════════════════════════════
        with tab1:
            
            # --- FORM THÊM CÁN BỘ DÀNH RIÊNG CHO ADMIN ---
            if st.session_state.role == "admin":
                with st.expander("➕ THÊM CÁN BỘ MỚI", expanded=False):
                    with st.form("form_add_nv", clear_on_submit=True):
                        st.markdown("<span style='color:#004B87; font-weight:bold;'>Điền thông tin cán bộ cần thêm:</span>", unsafe_allow_html=True)
                        c_add1, c_add2, c_add3, c_add4, c_add5 = st.columns(5)
                        n_hoten = c_add1.text_input("Họ và tên *")
                        n_chucvu = c_add2.text_input("Chức vụ")
                        n_mangach = c_add3.text_input("Mã ngạch")
                        n_bacluong = c_add4.text_input("Bậc lương")
                        n_heso = c_add5.text_input("Hệ số")
                        
                        n_ngay = st.text_input("Ngày nâng lương gần nhất (DD/MM/YYYY) *", placeholder="VD: 15/05/2021")
                        if st.form_submit_button("Lưu Cán bộ Mới", type="primary"):
                            if n_hoten and n_ngay:
                                supabase.table("theo_doi_luong").insert({
                                    "ho_ten": n_hoten, "chuc_vu": n_chucvu, 
                                    "ma_ngach": n_mangach, "bac_luong": n_bacluong,
                                    "he_so_hien_tai": n_heso, "ngay_gan_nhat": n_ngay
                                }).execute()
                                st.success("✅ Đã thêm thành công!")
                                st.rerun()
                            else:
                                st.error("⚠️ Vui lòng điền Tên và Ngày nâng lương!")
            
            # Filter bar
            with st.container():
                st.markdown('<div class="section-title">🔍 Bộ lọc tìm kiếm</div>', unsafe_allow_html=True)
                c1, c2, c3, c4, c5 = st.columns([2, 1.5, 1.8, 1, 1])
                search   = c1.text_input("Tên / chức vụ", placeholder="Gõ để tìm kiếm...", label_visibility="visible")
                
                loc_tg   = c2.selectbox("Trạng thái",
                    ["Tất cả", "Trong tháng này", "Trong Quý này", "Trong 6 tháng tới", "Trong năm nay", "Đã quá hạn"])
                loai_ngay = c3.selectbox("Loại ngày lọc",
                    ["Ngày dự kiến (Tương lai)", "Ngày gần nhất (Đã nâng)"])
                loc_nam  = c4.selectbox("Năm", ["Tất cả"] + [str(y) for y in range(2020, 2036)])
                loc_thang = c5.selectbox("Tháng", ["Tất cả"] + [str(m) for m in range(1, 13)])

            # Lọc dữ liệu theo đúng logic sếp yêu cầu
            df_display = df_calculated.copy()
            if search:
                df_display = df_display[df_display.astype(str).apply(lambda x: x.str.contains(search, case=False)).any(axis=1)]

            date_col = 'ngay_du_kien' if "dự kiến" in loai_ngay else 'ngay_gan_nhat'
            df_display['ngay_temp'] = pd.to_datetime(df_display[date_col], format='%d/%m/%Y', errors='coerce')
            today = datetime.now()

            if loc_tg != "Tất cả":
                if loc_tg == "Trong tháng này":
                    df_display = df_display[
                        (df_display['ngay_temp'].dt.month == today.month) &
                        (df_display['ngay_temp'].dt.year == today.year)]
                elif loc_tg == "Trong Quý này":
                    df_display = df_display[
                        (df_display['ngay_temp'].dt.quarter == (today.month-1)//3+1) &
                        (df_display['ngay_temp'].dt.year == today.year)]
                elif loc_tg == "Trong 6 tháng tới":
                    df_display = df_display[
                        (df_display['ngay_temp'].dt.date >= today.date()) &
                        (df_display['ngay_temp'].dt.date <= (today + relativedelta(months=6)).date())]
                elif loc_tg == "Trong năm nay":
                    df_display = df_display[df_display['ngay_temp'].dt.year == today.year]
                elif loc_tg == "Đã quá hạn":
                    df_display = df_display[df_display['ngay_temp'].dt.date < today.date()]

            # Bảo vệ chống lỗi khi thiếu ngày
            if loc_nam != "Tất cả":
                df_display = df_display[df_display['ngay_temp'].dt.year.fillna(-1).astype(int) == int(loc_nam)]
            if loc_thang != "Tất cả":
                df_display = df_display[df_display['ngay_temp'].dt.month.fillna(-1).astype(int) == int(loc_thang)]

            df_display['ma_ngach'] = df_display['ma_ngach'].apply(format_ma_ngach)
            df_display = df_display.drop(columns=['ngay_temp'], errors='ignore')

            st.markdown(f"<div style='font-size:12px; color:var(--muted); margin:8px 0;'>Hiển thị <b>{len(df_display)}</b> / {tong_nv} cán bộ</div>", unsafe_allow_html=True)

            # ── BẢNG DỮ LIỆU ────────────────────────
            disabled_cols = ["bac_luong_moi", "he_so_moi", "vuot_khung_moi", "ngay_du_kien", "trang_thai"]
            col_cfg = {"ma_ngach": st.column_config.TextColumn("Mã ngạch")}
            styled = df_display.style.map(style_trang_thai, subset=['trang_thai'])

            if st.session_state.role == "admin":
                st.info("💡 Chế độ Admin: Bạn có thể chỉnh sửa trực tiếp trên bảng dưới (Không thể thêm dòng ở đây, vui lòng dùng nút Thêm Cán Bộ Mới ở trên).")
                # Bỏ num_rows='dynamic' để ẩn nút thêm/xóa dòng đi
                edited_df = st.data_editor(
                    styled, use_container_width=True,
                    hide_index=True, column_config=col_cfg, disabled=disabled_cols
                )
                export_data = edited_df.data if hasattr(edited_df, 'data') else edited_df
            else:
                st.caption("👁 Chế độ xem: Bạn chỉ có thể xem bảng và biểu đồ phân tích. Liên hệ quản trị viên để cập nhật dữ liệu.")
                st.dataframe(styled, use_container_width=True, hide_index=True, column_config=col_cfg)
                export_data = df_display

            st.markdown("<hr class='divider'>", unsafe_allow_html=True)

            # ── ACTION BUTTONS ──────────────────────
            if st.session_state.role == "admin":
                cols = st.columns(3)
                col_luu = cols[0]
                col_excel = cols[1]
                col_word = cols[2]
                
                with col_luu:
                    if st.button("💾  Lưu thay đổi vào CSDL", use_container_width=True, type="primary"):
                        
                        # --- CƠ CHẾ BẢO VỆ DỮ LIỆU CHỐNG XÓA NHẦM ---
                        # Lấy lại bảng gốc, đắp những dòng vừa được sửa (export_data) lên bảng gốc
                        df_to_save = df_calculated.copy()
                        df_to_save.update(export_data)

                        recs = []
                        for r in df_to_save[df_to_save['ho_ten'].astype(str).str.strip().astype(bool)].to_dict(orient="records"):
                            recs.append({k: (None if pd.isna(v) or v == "" else v)
                                         for k, v in r.items()
                                         if k not in ["bac_luong_moi","he_so_moi","vuot_khung_moi","ngay_du_kien","trang_thai","id","ngay_temp"]})
                        
                        # Cập nhật an toàn lên Supabase
                        supabase.table("theo_doi_luong").delete().neq("ho_ten", "_PLACEHOLDER_").execute()
                        if recs:
                            supabase.table("theo_doi_luong").insert(recs).execute()
                        st.success("✅ Đã lưu dữ liệu thành công! (Dữ liệu của các cán bộ bị ẩn bởi bộ lọc vẫn được bảo toàn)")
                        st.rerun()
            else:
                cols = st.columns(2)
                col_excel = cols[0]
                col_word = cols[1]

            with col_excel:
                try:
                    buf_e = io.BytesIO()
                    with pd.ExcelWriter(buf_e, engine='openpyxl') as wr:
                        export_data.to_excel(wr, index=False, sheet_name='NangLuong')
                        ws = wr.sheets['NangLuong']
                        hdr_fill = PatternFill(start_color="0D47A1", end_color="0D47A1", fill_type="solid")
                        hdr_font = Font(bold=True, color="FFFFFF", name="Be Vietnam Pro", size=11)
                        for col_num in range(1, len(export_data.columns) + 1):
                            cell = ws.cell(row=1, column=col_num)
                            cell.fill = hdr_fill
                            cell.font = hdr_font
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                            ws.column_dimensions[get_column_letter(col_num)].width = 22
                        ws.row_dimensions[1].height = 28
                        stripe = PatternFill(start_color="EEF2FF", end_color="EEF2FF", fill_type="solid")
                        for row_idx in range(2, ws.max_row + 1):
                            if row_idx % 2 == 0:
                                for col_idx in range(1, len(export_data.columns) + 1):
                                    ws.cell(row=row_idx, column=col_idx).fill = stripe
                    st.download_button(
                        "📥  Xuất báo cáo Excel", buf_e.getvalue(),
                        file_name=f"NangLuong_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                except Exception as ex:
                    st.warning(f"Lỗi xuất Excel: {ex}")

            with col_word:
                st.download_button(
                    "📝  Xuất tờ trình Word", tao_file_word_dien_bien(export_data, loc_thang, loc_nam),
                    file_name=f"DienBienLuong_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        # ══════════════════════════════════════════
        # TAB 2 – DASHBOARD
        # ══════════════════════════════════════════
        with tab2:
            CHART_H = 300
            MARGINS = dict(t=40, b=10, l=10, r=10)
            TITLE_FONT = dict(size=13, color='#0d47a1', family='Be Vietnam Pro')
            PLOT_BG   = 'rgba(0,0,0,0)'
            GRID_COL  = '#e8ecf5'

            r2c1, r2c2 = st.columns(2)

            df_ma = df_calculated['ma_ngach'].value_counts().reset_index()
            df_ma = df_ma[df_ma['ma_ngach'].astype(str).str.strip() != ""]
            fig_ma = px.bar(df_ma, x='ma_ngach', y='count', text='count',
                            color='count', color_continuous_scale='Blues')
            fig_ma.update_traces(textposition='outside', marker_line_color='#0d47a1',
                                 marker_line_width=1, opacity=.88, textfont_size=12)
            fig_ma.update_layout(
                title=dict(text="Phân bố Mã ngạch", x=.5, font=TITLE_FONT),
                xaxis_title="", yaxis_title="", coloraxis_showscale=False,
                plot_bgcolor=PLOT_BG, paper_bgcolor=PLOT_BG,
                height=CHART_H, margin=MARGINS, font_family="Be Vietnam Pro"
            )
            fig_ma.update_yaxes(showgrid=True, gridcolor=GRID_COL, showticklabels=False)
            with r2c1:
                st.markdown('<div class="chart-wrap">', unsafe_allow_html=True)
                st.plotly_chart(fig_ma, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)

            df_ngach = df_calculated['ngach_luong'].value_counts().reset_index()
            df_ngach = df_ngach[df_ngach['ngach_luong'].astype(str).str.strip() != ""]
            fig_ngach = px.bar(df_ngach, y='ngach_luong', x='count', text='count',
                               orientation='h', color='count', color_continuous_scale='Teal')
            fig_ngach.update_traces(textposition='outside', marker_line_color='#00838f',
                                    marker_line_width=1, opacity=.88, textfont_size=12)
            fig_ngach.update_layout(
                title=dict(text="Phân bố Ngạch lương", x=.5, font=TITLE_FONT),
                xaxis_title="", yaxis_title="", coloraxis_showscale=False,
                plot_bgcolor=PLOT_BG, paper_bgcolor=PLOT_BG,
                height=CHART_H, margin=MARGINS, font_family="Be Vietnam Pro"
            )
            fig_ngach.update_xaxes(showgrid=True, gridcolor=GRID_COL, showticklabels=False)
            with r2c2:
                st.markdown('<div class="chart-wrap">', unsafe_allow_html=True)
                st.plotly_chart(fig_ngach, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)

            r3c1, r3c2 = st.columns(2)

            df_bac = df_calculated['bac_luong'].value_counts().reset_index()
            df_bac = df_bac[df_bac['bac_luong'].astype(str).str.strip() != ""]
            fig_bac = px.pie(df_bac, names='bac_luong', values='count', hole=.55,
                             color_discrete_sequence=px.colors.sequential.Blues_r)
            fig_bac.update_traces(textposition='inside', textinfo='percent+label',
                                  marker=dict(line=dict(color='#fff', width=2)))
            fig_bac.update_layout(
                title=dict(text="Bậc lương hiện hưởng", x=.5, font=TITLE_FONT),
                showlegend=False, height=CHART_H, margin=MARGINS,
                paper_bgcolor=PLOT_BG, font_family="Be Vietnam Pro"
            )
            with r3c1:
                st.markdown('<div class="chart-wrap">', unsafe_allow_html=True)
                st.plotly_chart(fig_bac, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)

            df_tt = df_calculated['trang_thai'].value_counts().reset_index()
            df_tt = df_tt[df_tt['trang_thai'].astype(str).str.strip() != ""]
            color_map = {
                "🟢 Chưa đến hạn": "#2e7d32",
                "🟡 Sắp đến hạn (Quý này)": "#f57f17",
                "🔴 Sắp đến hạn (Tháng này)": "#e65100",
                "⛔ Đã quá hạn": "#b71c1c",
                "Chưa có ngày": "#9e9e9e",
            }
            colors = [color_map.get(v, "#78909c") for v in df_tt['trang_thai']]
            fig_tt = go.Figure(go.Bar(
                x=df_tt['count'], y=df_tt['trang_thai'], orientation='h',
                text=df_tt['count'], textposition='outside',
                marker_color=colors, marker_line_width=0,
            ))
            fig_tt.update_layout(
                title=dict(text="Trạng thái nâng lương", x=.5, font=TITLE_FONT),
                xaxis_title="", yaxis_title="", plot_bgcolor=PLOT_BG,
                paper_bgcolor=PLOT_BG, height=CHART_H, margin=MARGINS,
                font_family="Be Vietnam Pro"
            )
            fig_tt.update_xaxes(showgrid=True, gridcolor=GRID_COL, showticklabels=False)
            with r3c2:
                st.markdown('<div class="chart-wrap">', unsafe_allow_html=True)
                st.plotly_chart(fig_tt, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)

        # ══════════════════════════════════════════
        # TAB 3 – VĂN BẢN NÂNG LƯƠNG (CHỈ ADMIN)
        # ══════════════════════════════════════════
        if st.session_state.role == "admin":
            with tab3:
                st.markdown('<div class="section-title">🏛️ Xây dựng Quyết định / Tờ trình / Biên bản nâng lương</div>', unsafe_allow_html=True)
                st.caption("Chọn các cán bộ đến đợt xét nâng lương trong kỳ này. Hệ thống tự động phân loại: "
                           "cán bộ **Trưởng ban / Phó Trưởng ban** → **Tờ trình gửi Ban Tổ chức Tỉnh ủy**; "
                           "các cán bộ, chuyên viên khác → **Quyết định** do Trưởng Ban ký ban hành trực tiếp. "
                           "Cả hai đều kèm theo **Biên bản họp Hội đồng xét nâng lương**.")

                ds_ten = df_calculated['ho_ten'].astype(str).tolist()
                chon_ten = st.multiselect("👥 Chọn cán bộ đến đợt xét nâng lương", ds_ten)

                if chon_ten:
                    ds_rows_full = []
                    for ten in chon_ten:
                        row = df_calculated[df_calculated['ho_ten'] == ten].iloc[0].to_dict()
                        row['loai'] = xac_dinh_loai_nang_luong(row)
                        row['la_lanh_dao'] = la_lanh_dao(row)
                        ds_rows_full.append(row)

                    nhom_qd_tx = [r for r in ds_rows_full if not r['la_lanh_dao'] and r['loai'] == 'thuong_xuyen']
                    nhom_qd_vk = [r for r in ds_rows_full if not r['la_lanh_dao'] and r['loai'] == 'vuot_khung']
                    nhom_tt_tx = [r for r in ds_rows_full if r['la_lanh_dao'] and r['loai'] == 'thuong_xuyen']
                    nhom_tt_vk = [r for r in ds_rows_full if r['la_lanh_dao'] and r['loai'] == 'vuot_khung']

                    def _ten_list(rows):
                        return ", ".join(r['ho_ten'] for r in rows) if rows else "—"

                    st.markdown(f"""
                    <div class="vb-group">
                    <b>📄 Quyết định — nâng lương thường xuyên</b> ({len(nhom_qd_tx)}): {_ten_list(nhom_qd_tx)}<br>
                    <b>📄 Quyết định — nâng phụ cấp vượt khung</b> ({len(nhom_qd_vk)}): {_ten_list(nhom_qd_vk)}<br>
                    <b>📝 Tờ trình BTC Tỉnh ủy — nâng lương thường xuyên</b> ({len(nhom_tt_tx)}): {_ten_list(nhom_tt_tx)}<br>
                    <b>📝 Tờ trình BTC Tỉnh ủy — nâng phụ cấp vượt khung</b> ({len(nhom_tt_vk)}): {_ten_list(nhom_tt_vk)}
                    </div>
                    """, unsafe_allow_html=True)

                    st.markdown("##### ⚙️ Thông tin chung của đợt xét nâng lương")
                    cA, cB, cC = st.columns(3)
                    so_qd = cA.text_input("Số Quyết định (không gồm '-QĐ/BTGDVTU')", value="")
                    so_tt = cB.text_input("Số Tờ trình (không gồm '-TTr/BTGDV')", value="")
                    ngay_hop_bb = cC.text_input("Ngày họp Hội đồng (dd/mm/yyyy)", value=datetime.now().strftime("%d/%m/%Y"))

                    cD, cE, cF = st.columns(3)
                    ngay_ky = cD.text_input("Ngày ký văn bản", value=datetime.now().strftime("%d"))
                    thang_ky = cE.text_input("Tháng ký văn bản", value=datetime.now().strftime("%m"))
                    nam_ky = cF.text_input("Năm ký văn bản", value=datetime.now().strftime("%Y"))

                    cG, cH = st.columns(2)
                    gio_bat_dau = cG.text_input("Giờ họp bắt đầu", value="08 giờ 00")
                    gio_ket_thuc = cH.text_input("Giờ họp kết thúc", value="09 giờ 00")

                    cI, cJ = st.columns(2)
                    truong_ban = cI.text_input("Trưởng Ban (người ký QĐ/Tờ trình, chủ trì họp)", value="Trần Mạnh Lợi")
                    thu_ky = cJ.text_input("Thư ký Hội đồng", value="Đinh Thị Thúy")

                    thanh_phan_text = st.text_area("👥 Thành phần Hội đồng xét nâng lương (mỗi người 1 dòng)",
                                                     value=THANH_PHAN_HOI_DONG_MAC_DINH, height=200)

                    st.markdown("<hr class='divider'>", unsafe_allow_html=True)

                    st.caption("💡 Nếu Ban ký số/cấp số điện tử, có thể để trống Số Quyết định / Số Tờ trình — "
                               "văn bản sẽ để chỗ trống (Số ......../...) để hệ thống ký số điền sau.")

                    if st.button("📄  Tạo văn bản", type="primary", use_container_width=True):
                        if True:
                            files = {}
                            if nhom_qd_tx:
                                files[f"QuyetDinh_ThuongXuyen_{datetime.now().strftime('%Y%m%d')}.docx"] = tao_quyet_dinh(
                                    nhom_qd_tx, "thuong_xuyen", so_qd, ngay_ky, thang_ky, nam_ky, ngay_hop_bb, truong_ban)
                                files[f"BienBan_ThuongXuyen_CBCC_{datetime.now().strftime('%Y%m%d')}.docx"] = tao_bien_ban(
                                    nhom_qd_tx, "thuong_xuyen", False, ngay_hop_bb, gio_bat_dau, gio_ket_thuc,
                                    thanh_phan_text, truong_ban=truong_ban, thu_ky=thu_ky)
                            if nhom_qd_vk:
                                files[f"QuyetDinh_VuotKhung_{datetime.now().strftime('%Y%m%d')}.docx"] = tao_quyet_dinh(
                                    nhom_qd_vk, "vuot_khung", so_qd, ngay_ky, thang_ky, nam_ky, ngay_hop_bb, truong_ban)
                                files[f"BienBan_VuotKhung_CBCC_{datetime.now().strftime('%Y%m%d')}.docx"] = tao_bien_ban(
                                    nhom_qd_vk, "vuot_khung", False, ngay_hop_bb, gio_bat_dau, gio_ket_thuc,
                                    thanh_phan_text, truong_ban=truong_ban, thu_ky=thu_ky)
                            if nhom_tt_tx:
                                files[f"ToTrinh_ThuongXuyen_{datetime.now().strftime('%Y%m%d')}.docx"] = tao_to_trinh(
                                    nhom_tt_tx, "thuong_xuyen", so_tt, ngay_ky, thang_ky, nam_ky, ngay_hop_bb, truong_ban)
                                files[f"BienBan_ThuongXuyen_LanhDao_{datetime.now().strftime('%Y%m%d')}.docx"] = tao_bien_ban(
                                    nhom_tt_tx, "thuong_xuyen", True, ngay_hop_bb, gio_bat_dau, gio_ket_thuc,
                                    thanh_phan_text, truong_ban=truong_ban, thu_ky=thu_ky)
                            if nhom_tt_vk:
                                files[f"ToTrinh_VuotKhung_{datetime.now().strftime('%Y%m%d')}.docx"] = tao_to_trinh(
                                    nhom_tt_vk, "vuot_khung", so_tt, ngay_ky, thang_ky, nam_ky, ngay_hop_bb, truong_ban)
                                files[f"BienBan_VuotKhung_LanhDao_{datetime.now().strftime('%Y%m%d')}.docx"] = tao_bien_ban(
                                    nhom_tt_vk, "vuot_khung", True, ngay_hop_bb, gio_bat_dau, gio_ket_thuc,
                                    thanh_phan_text, truong_ban=truong_ban, thu_ky=thu_ky)

                            st.success(f"✅ Đã tạo {len(files)} văn bản. Tải về bên dưới:")
                            cols_dl = st.columns(len(files))
                            for i, (fname, fdata) in enumerate(files.items()):
                                with cols_dl[i]:
                                    st.download_button(
                                        f"📥 {fname}", fdata, file_name=fname,
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        use_container_width=True, key=f"dl_{fname}"
                                    )
                else:
                    st.info("👆 Chọn ít nhất một cán bộ ở trên để bắt đầu.")

    except Exception as e:
        st.error(f"Lỗi hệ thống: {e}")
        import traceback
        st.code(traceback.format_exc())

if __name__ == "__main__":
    main()
